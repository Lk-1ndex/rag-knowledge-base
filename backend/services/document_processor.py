import json
import re
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document as DocxDocument

from config import settings


@dataclass
class TextPage:
    text: str
    page_number: int | None = None


@dataclass
class Block:
    """文档的一个语义块。heading_level=0 表示正文，1/2/3 表示标题层级。

    标题感知切分依赖这个结构：提取阶段把文档拆成有层级标记的块，
    切分阶段在标题边界强制断开，并把标题路径作为前缀注入正文 chunk。
    """

    text: str
    heading_level: int = 0
    page_number: int | None = None


@dataclass
class Chunk:
    text: str
    chunk_index: int
    page_number: int | None = None


SUPPORTED_TYPES = {".pdf": "pdf", ".docx": "docx", ".md": "md", ".txt": "txt"}


def file_type_from_name(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_TYPES:
        raise ValueError("仅支持 PDF、DOCX、Markdown、TXT 文件")
    return SUPPORTED_TYPES[suffix]


def secure_filename(filename: str) -> str:
    name = Path(filename).name
    name = re.sub(r"[^\w.\-\u4e00-\u9fa5]+", "_", name, flags=re.UNICODE).strip("._")
    return name or "upload.bin"


def validate_magic_bytes(filename: str, data: bytes) -> None:
    file_type = file_type_from_name(filename)
    head = data[:8]
    if file_type == "pdf" and not head.startswith(b"%PDF"):
        raise ValueError("PDF 文件头校验失败")
    if file_type == "docx" and not head.startswith(b"PK"):
        raise ValueError("DOCX 文件头校验失败")
    if file_type in {"md", "txt"}:
        try:
            data[:4096].decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("文本文件必须使用 UTF-8 编码") from exc


def extract_text_pages(path: Path, file_type: str) -> list[TextPage]:
    if file_type == "pdf":
        return extract_pdf(path)
    if file_type == "docx":
        return extract_docx(path)
    return [TextPage(path.read_text(encoding="utf-8"), None)]


def extract_blocks(path: Path, file_type: str) -> list[Block]:
    """提取带标题层级的块列表，供标题感知切分使用。"""
    if file_type == "pdf":
        return extract_pdf_blocks(path)
    if file_type == "docx":
        return extract_docx_blocks(path)
    if file_type == "md":
        return extract_markdown_blocks(path.read_text(encoding="utf-8"))
    return [Block(text=path.read_text(encoding="utf-8"), heading_level=0, page_number=None)]


_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def extract_markdown_blocks(text: str) -> list[Block]:
    """按 Markdown 标题（# / ## / ###）拆分，标题行单独成块并标记层级。"""
    blocks: list[Block] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            body = "\n".join(buffer).strip()
            if body:
                blocks.append(Block(text=body, heading_level=0))
            buffer.clear()

    for line in text.splitlines():
        match = _MD_HEADING.match(line.strip())
        if match:
            flush()
            level = min(len(match.group(1)), 3)
            title = match.group(2).strip()
            if title:
                blocks.append(Block(text=title, heading_level=level))
        else:
            buffer.append(line)
    flush()
    return blocks or [Block(text=text.strip(), heading_level=0)]


def extract_docx_blocks(path: Path) -> list[Block]:
    """DOCX 用段落 style 名识别标题（Heading 1/2/3 或中文'标题 1'）。"""
    doc = DocxDocument(path)
    blocks: list[Block] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = _docx_heading_level(para)
        blocks.append(Block(text=text, heading_level=level))
    return blocks or [Block(text="", heading_level=0)]


def _docx_heading_level(para) -> int:
    style_name = (getattr(para.style, "name", "") or "").lower()
    match = re.search(r"(?:heading|标题)\s*([1-3])", style_name)
    if match:
        return int(match.group(1))
    return 0


def extract_pdf_blocks(path: Path) -> list[Block]:
    """PDF 启发式标题识别：用字号信息，明显大于正文中位数字号或加粗的短行判为标题。"""
    raw_lines: list[tuple[str, float, bool, int]] = []  # (text, size, bold, page)
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = "".join(span.get("text", "") for span in spans).strip()
                    if not line_text:
                        continue
                    max_size = max(float(span.get("size", 0)) for span in spans)
                    is_bold = any(int(span.get("flags", 0)) & 16 for span in spans)
                    raw_lines.append((line_text, max_size, is_bold, index))

    if not raw_lines:
        return [Block(text="", heading_level=0)]

    body_size = _median([size for _, size, _, _ in raw_lines])
    blocks: list[Block] = []
    for text, size, is_bold, page in raw_lines:
        level = _pdf_heading_level(text, size, is_bold, body_size)
        blocks.append(Block(text=text, heading_level=level, page_number=page))
    return blocks


def _pdf_heading_level(text: str, size: float, is_bold: bool, body_size: float) -> int:
    """启发式：短行 + 字号显著大于正文 → 标题。字号越大层级越高。"""
    if len(text) > 80:  # 长行几乎不可能是标题
        return 0
    ratio = size / body_size if body_size else 1.0
    if ratio >= 1.5:
        return 1
    if ratio >= 1.25:
        return 2
    if ratio >= 1.1 or (is_bold and ratio >= 1.05):
        return 3
    return 0


def _median(values: list[float]) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    mid = len(ordered) // 2
    if len(ordered) % 2:
        return ordered[mid]
    return (ordered[mid - 1] + ordered[mid]) / 2


def extract_pdf(path: Path) -> list[TextPage]:
    pages: list[TextPage] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append(TextPage(text=text, page_number=index))
    return pages


def extract_docx(path: Path) -> list[TextPage]:
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return [TextPage("\n\n".join(paragraphs), None)]


def rough_token_len(text: str) -> int:
    ascii_words = re.findall(r"[A-Za-z0-9_]+", text)
    cjk_chars = re.findall(r"[\u4e00-\u9fff]", text)
    other = max(len(text) - sum(len(w) for w in ascii_words) - len(cjk_chars), 0) // 4
    return len(ascii_words) + len(cjk_chars) + other


def split_paragraphs(text: str) -> list[str]:
    parts = re.split(r"\n\s*\n", text)
    return [part.strip() for part in parts if part.strip()]


def chunk_blocks(
    blocks: list[Block],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Chunk]:
    """标题感知切分：在标题边界强制断开，把标题路径作为前缀注入每个 chunk。

    维护一个标题路径栈（如 ["方法", "信道估计"]），遇到标题时按层级更新栈，
    遇到正文时按 chunk_size 切分并在每个 chunk 前加上 "方法 > 信道估计" 前缀，
    让检索时能命中"这一段属于哪个章节"的语义信息。
    """
    size = chunk_size or settings.default_chunk_size
    overlap = chunk_overlap or settings.default_chunk_overlap
    chunks: list[Chunk] = []
    heading_stack: list[str] = []

    def prefix() -> str:
        return " > ".join(heading_stack)

    def emit(body: str, page: int | None) -> None:
        if not body.strip():
            return
        path = prefix()
        text = f"[{path}]\n{body}" if path else body
        chunks.append(Chunk(text=text, chunk_index=len(chunks), page_number=page))

    current = ""
    current_page: int | None = None

    def flush() -> None:
        nonlocal current
        if current:
            emit(current, current_page)
            current = ""

    for block in blocks:
        if block.heading_level > 0:
            # 标题边界：先冲刷正文缓冲，再更新标题路径栈
            flush()
            level = block.heading_level
            del heading_stack[level - 1 :]
            heading_stack.append(block.text)
            continue

        current_page = block.page_number
        for paragraph in split_paragraphs(block.text):
            candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
            if rough_token_len(candidate) <= size:
                current = candidate
                continue
            if current:
                emit(current, current_page)
                current = build_overlap(current, overlap)
            if rough_token_len(paragraph) > size:
                for piece in split_long_text(paragraph, size):
                    candidate = f"{current}\n\n{piece}".strip() if current else piece
                    emit(candidate, current_page)
                    current = build_overlap(candidate, overlap)
            else:
                current = f"{current}\n\n{paragraph}".strip() if current else paragraph

    flush()
    return chunks


def chunk_pages(
    pages: list[TextPage],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Chunk]:
    size = chunk_size or settings.default_chunk_size
    overlap = chunk_overlap or settings.default_chunk_overlap
    chunks: list[Chunk] = []
    carry = ""

    for page in pages:
        current = carry
        for paragraph in split_paragraphs(page.text):
            candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
            if rough_token_len(candidate) <= size:
                current = candidate
                continue
            if current:
                chunks.append(Chunk(current, len(chunks), page.page_number))
                current = build_overlap(current, overlap)
            if rough_token_len(paragraph) > size:
                for piece in split_long_text(paragraph, size):
                    candidate = f"{current}\n\n{piece}".strip() if current else piece
                    chunks.append(Chunk(candidate, len(chunks), page.page_number))
                    current = build_overlap(candidate, overlap)
            else:
                current = f"{current}\n\n{paragraph}".strip() if current else paragraph
        carry = current

    if carry:
        chunks.append(Chunk(carry, len(chunks), pages[-1].page_number if pages else None))
    return chunks


def split_long_text(text: str, size: int) -> list[str]:
    sentences = re.split(r"(?<=[。！？.!?])\s*", text)
    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        if not sentence:
            continue
        candidate = current + sentence
        if rough_token_len(candidate) <= size:
            current = candidate
        else:
            if current:
                pieces.append(current)
            current = sentence
    if current:
        pieces.append(current)
    return pieces


def build_overlap(text: str, overlap: int) -> str:
    if overlap <= 0:
        return ""
    words = re.findall(r"\S+", text)
    return " ".join(words[-overlap:])


def parse_tags(raw: str | None) -> str:
    if not raw:
        return "[]"
    tags = [tag.strip() for tag in raw.split(",") if tag.strip()]
    return json.dumps(tags, ensure_ascii=False)
